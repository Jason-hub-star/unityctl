#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""2026 오픈소스 개발자대회 결과보고서(.docx) 자동 채움 — unityctl.
   팀=URHYNIX(2명, 일반) / AI 티 제거 윤문 / 안내 페이지 삭제 / SBOM 보강."""
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = sys.argv[1]
OUT = sys.argv[2]
FONT = "맑은 고딕"


def style_run(run):
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(a), FONT)


def set_cell(cell, text):
    paras = cell.paragraphs
    first = paras[0]
    for r in list(first.runs):
        r.text = ""
    for p in paras[1:]:
        p._element.getparent().remove(p._element)
    lines = str(text).split("\n")
    run = first.add_run(lines[0])
    style_run(run)
    for line in lines[1:]:
        p = cell.add_paragraph()
        run = p.add_run(line)
        style_run(run)


def delete_guide_page(doc):
    """본문 제목 표(T1) 앞의 모든 안내 요소 제거."""
    body = doc.element.body
    title = doc.tables[1]._element
    for child in list(body):
        if child is title:
            break
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)
    # 제목 단락의 pageBreakBefore 제거(빈 첫 장 방지)
    for p in doc.paragraphs[:3]:
        pPr = p._p.find(qn("w:pPr"))
        if pPr is not None:
            pb = pPr.find(qn("w:pageBreakBefore"))
            if pb is not None:
                pPr.remove(pb)


doc = Document(SRC)
T = doc.tables

# ── Table 2: 기본정보 ─────────────────────────────
set_cell(T[2].rows[1].cells[1], "URHYNIX")          # 팀명
set_cell(T[2].rows[1].cells[3], "2명")               # 팀 인원(팀장 포함)
set_cell(T[2].rows[2].cells[1], "일반")              # 참가부문
set_cell(T[2].rows[2].cells[3], "자유과제")          # 과제유형

# ── Table 3: 프로젝트 개요/세부 ───────────────────
set_cell(T[3].rows[1].cells[1], "unityctl - AI 에이전트를 위한 Unity 에디터 제어 도구")
set_cell(T[3].rows[2].cells[1], "https://github.com/Jason-hub-star/unityctl")
set_cell(T[3].rows[3].cells[1], "[유튜브 업로드 후 URL 기재 — 제출 전 작성]")

set_cell(T[3].rows[4].cells[1],
    "AI 에이전트가 Unity 에디터를 직접 조작하도록 연결해 주는 오픈소스 CLI이자 MCP 서버입니다. "
    "씬과 게임오브젝트를 다루고 플레이모드를 돌리는 일부터 스크린샷으로 결과를 확인하는 것까지, "
    "170개가 넘는 명령을 명령줄과 CI에서 그대로 쓸 수 있습니다.")

set_cell(T[3].rows[6].cells[1],
    "- 요즘 AI 코딩 도구는 웹이나 앱은 곧잘 짜주지만, 정작 게임을 만드는 Unity 에디터 안으로는 손을 뻗지 못합니다.\n"
    "- 2026년 Unity가 'Unity AI'를 내놨지만, 어디까지나 에디터 안에서 사람이 대화로 쓰는 보조 기능이고 특정 구독에 묶여 있습니다.\n"
    "- 외부의 AI 에이전트가 명령줄이나 CI에서 에디터를 직접, 그것도 일관된 결과로 제어할 방법은 여전히 비어 있었습니다.\n"
    "- unityctl은 바로 그 자리를 채웁니다. 어떤 LLM이든 자연어로 Unity를 부리고 결과까지 스스로 확인하게 하는, 특정 회사에 묶이지 않는 오픈소스를 목표로 만들었습니다.")

set_cell(T[3].rows[7].cells[1],
    "- 런타임: .NET 10 (SDK 10.0.201), Shared 모듈은 .NET Standard 2.1\n"
    "- 언어: C# (CLI·Core·MCP)와 Unity Editor 플러그인(UPM)\n"
    "- CLI 프레임워크: ConsoleAppFramework 5.3.3, 출력 렌더링은 Spectre.Console\n"
    "- AI 연동: ModelContextProtocol 1.1.0 기반 MCP 서버\n"
    "- 직렬화: System.Text.Json(CLI/Core), Newtonsoft.Json(Plugin), 스크린샷 처리는 SixLabors.ImageSharp\n"
    "- 테스트와 CI: xUnit 887개, GitHub Actions 3-OS 매트릭스(Windows·macOS·Linux)\n"
    "- 지원 Unity 버전: 2021.3 이상")

set_cell(T[3].rows[8].cells[1],
    "- 계층은 Shared(프로토콜·모델)를 Core(전송·탐색·재시도)가 받고, 그 위에 얇은 CLI를 얹는 구조입니다. MCP 서버는 별도로 둡니다.\n"
    "- Unity 플러그인은 에디터 안에서 도는 IPC 서버 역할을 합니다.\n"
    "- 전송은 IPC(Named Pipe)를 먼저 찔러보고, 안 되면 batchmode 실행으로 자동으로 넘어갑니다.\n"
    "- IPC 프레이밍은 [4바이트 LE 길이][UTF-8 JSON] 형태로, 최대 10MB까지 받습니다.\n"
    "- 도메인 리로드로 연결이 끊기는 구간은 ipc-state.json 하트비트와 reload 인지 재시도로 메웁니다.\n"
    "- screenshot과 workflow verify로 에이전트가 자기 작업 결과를 직접 눈으로 확인하는 폐루프를 갖췄습니다.")

set_cell(T[3].rows[9].cells[1],
    "프로젝트 상세 내용\n"
    "- 170개가 넘는 CLI 명령: 씬 계층과 게임오브젝트·컴포넌트 조작, 플레이모드, 스크린샷, "
    "스크립트 생성·편집·패치, 프로파일러, 머티리얼·애니메이션·UI(UGUI/UI Toolkit)·Cinemachine까지 다룹니다.\n"
    "- MCP 하이브리드: 통합 도구 몇 개로 수많은 명령을 토큰 효율적으로 노출합니다(33개를 12개로 묶음).\n"
    "- 멀티 에디터 라우팅: editor current/select로 여러 인스턴스를 골라 보낼 수 있습니다.\n"
    "- 운영 기능: 비동기 명령, 세션 레이어, Flight Recorder, Ghost Mode(--dry-run), Watch Mode를 갖췄습니다.\n"
    "- 시각 검증 워크플로우: workflow verify와 playSmoke로 게임이 실제로 도는 장면을 캡처해 판정합니다.\n\n"
    "구동 및 시연\n"
    "- 설치: NuGet으로 CLI를 받고, Unity 프로젝트에 UPM 플러그인을 추가합니다.\n"
    "- 실행: dotnet run --project src/Unityctl.Cli -- <command>로 직접 쓰거나, MCP 서버로 Claude·Cursor 등과 연결합니다.\n"
    "- 검증: dotnet test 887개가 통과하고 3-OS CI가 초록불을 유지합니다.")

set_cell(T[3].rows[10].cells[1],
    "향후 확장성 및 기대효과\n"
    "- 반복되는 씬 구성과 컴포넌트 설정, 빌드 점검을 자연어 한 줄로 대신하니 개발 속도가 붙습니다.\n"
    "- Unity API를 깊이 몰라도 말로 프로토타입을 만들 수 있어, 1인 개발자나 교육 현장에서 특히 쓸모가 큽니다.\n"
    "- AI와 게임엔진을 잇는 표준 MCP 인터페이스를 887개 테스트와 CI까지 통째로 공개합니다.\n"
    "- Claude든 Cursor든 가리지 않고 붙고, 기존 CI/CD 파이프라인에도 그대로 얹힙니다.\n"
    "- 스크린샷으로 결과를 스스로 확인하는 구조라, 코드만 뱉고 끝나는 게 아니라 실제로 돌아가는 게임까지 끌고 갑니다.")

set_cell(T[3].rows[11].cells[1],
    "프로젝트의 혁신성 및 차별성\n"
    "- Unity AI는 에디터 창 안에서 대화로만 쓰고 구독이 필요하지만, unityctl은 외부 AI가 헤드리스 환경이나 CI에서 에디터를 일관되게 제어합니다.\n"
    "- MIT 라이선스 오픈소스라 특정 LLM이나 구독에 발이 묶이지 않습니다.\n"
    "- CLI와 MCP를 분리했고, IPC를 먼저 시도하다 안 되면 batchmode로 알아서 떨어집니다.\n"
    "- 스크린샷 기반 시각 검증과 여러 에디터 인스턴스 라우팅을 함께 갖췄습니다.\n"
    "- 887개 테스트와 3개 OS CI로 완성도를 숫자로 보여줍니다.\n\n"
    "한계점 및 향후 발전 로드맵\n"
    "- 아직 UI Toolkit처럼 손이 덜 닿은 도메인이 남아 있어 명령 커버리지를 더 넓혀야 합니다.\n"
    "- Unity 버전과 렌더 파이프라인별 호환성은 추가 검증이 필요합니다.\n"
    "- 앞으로는 명령 범위 확대, 시각 검증 고도화, 설치·문서 온보딩 개선을 차례로 풀어갈 계획입니다.\n\n"
    "소감 및 후기\n"
    "- 'AI한테 게임 만드는 손을 쥐여주자'는 한 줄 아이디어로 시작했는데, 막상 Unity 에디터를 외부에서 안정적으로 붙드는 일이 생각보다 험난했습니다. "
    "도메인 리로드 한 번에 연결이 통째로 끊기던 문제를 하트비트 상태 파일로 잡아낸 순간이 가장 기억에 남습니다.\n"
    "- 둘이서 한 명은 Core 전송 계층을, 한 명은 Unity 플러그인을 맡아 IPC 양쪽을 맞춰가는 과정이 곧 협업 그 자체였습니다. "
    "테스트를 887개까지 쌓아두니 리팩터링이 두렵지 않았고, 그 안정감이 끝까지 밀어붙이는 힘이 됐습니다.\n"
    "- 오픈소스로 공개하며 받은 피드백 하나하나가 방향을 잡아줬습니다. 이번 대회를 계기로 AI와 게임 개발이 만나는 길을 더 많은 분과 함께 넓혀가고 싶습니다.")

# ── Table 5: SBOM (의존성 6개) ───────────────────
SBOM = [
    ("1", "ConsoleAppFramework", "5.3.3", "MIT",
     "https://github.com/Cysharp/ConsoleAppFramework", "CLI 명령 등록·라우팅 프레임워크"),
    ("2", "ModelContextProtocol", "1.1.0", "MIT",
     "https://github.com/modelcontextprotocol/csharp-sdk", "MCP 서버 구현(AI 에이전트 연동)"),
    ("3", "Spectre.Console", "0.49", "MIT",
     "https://github.com/spectreconsole/spectre.console", "CLI 출력·테이블 렌더링"),
    ("4", "SixLabors.ImageSharp", "3.1.12", "Six Labors Split License (오픈소스는 Apache-2.0)",
     "https://github.com/SixLabors/ImageSharp", "스크린샷 이미지 인코딩·처리"),
    ("5", "System.Text.Json", "8.0.5", "MIT",
     "https://github.com/dotnet/runtime", "CLI/Core JSON 직렬화(Source Generator)"),
    ("6", "Newtonsoft.Json (Unity)", "3.2.1", "MIT",
     "https://github.com/JamesNK/Newtonsoft.Json", "Unity 플러그인 JSON 직렬화"),
]
for i, row in enumerate(SBOM):
    cells = T[5].rows[i + 1].cells
    for j, val in enumerate(row):
        set_cell(cells[j], val)

# ── Table 8: AI 모델 활용 명세 (unityctl은 AI 모델 미탑재) ──
set_cell(T[8].rows[3].cells[1], "해당 없음 (AI 모델 미탑재)")
set_cell(T[8].rows[3].cells[4], "해당 없음")
for r in (5, 6, 7, 8):
    set_cell(T[8].rows[r].cells[1], "해당 없음 (외부·자체 AI 모델 미사용)")
set_cell(T[8].rows[10].cells[1], "MIT License (OSI 인증)")
set_cell(T[8].rows[10].cells[4], "해당 없음 (AI 모델 학습·추론 코드 없음)")
set_cell(T[8].rows[11].cells[1],
    "코드 작성과 디버깅 보조용으로 Anthropic Claude, OpenAI GPT 등 상용 AI를 활용했습니다. "
    "다만 프로젝트에 AI 모델을 탑재하거나 학습하지 않으므로 유형 1·2·3에는 해당하지 않습니다.")

# ── 안내 페이지 삭제 ──────────────────────────────
delete_guide_page(doc)

doc.save(OUT)
print("saved:", OUT)
